from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_DIR = Path(r"C:\DATN\Word\UseCase\UseCaseSpecification")
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)


def set_run_font(run, bold=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = FONT_SIZE
    if bold is not None:
        run.bold = bold


def apply_doc_defaults(doc):
    styles = doc.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = FONT_NAME
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            style.font.size = FONT_SIZE


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_table(table, widths=None):
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            if widths and idx < len(widths):
                cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                if row_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, bold=True if row_idx == 0 else None)


def add_para(doc, text="", bold=False, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    return paragraph


def add_section_title(doc, text):
    return add_para(doc, text, bold=True)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    format_table(table, widths)
    add_para(doc)
    return table


def build_doc(usecase):
    doc = Document()
    apply_doc_defaults(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_para(doc, usecase["title"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_section_title(doc, "1. Mã use case (Use case code)")
    add_para(doc, usecase["code"])

    add_section_title(doc, "2. Mô tả ngắn (Brief description)")
    add_para(doc, usecase["description"])

    add_section_title(doc, "3. Actors")
    for idx, actor in enumerate(usecase["actors"], start=1):
        add_para(doc, f"3.{idx} Name of Actor {idx}: {actor}")

    add_section_title(doc, "4. Preconditions")
    for item in usecase["preconditions"]:
        add_para(doc, item)

    add_section_title(doc, "5. Basic Flow of Events")
    for step in usecase["basic_flow"]:
        add_para(doc, step)

    add_section_title(doc, "6. Alternative flows")
    add_table(
        doc,
        ["No", "Location", "Condition", "Action", "Resume location"],
        usecase["alternative_flows"],
        [Inches(0.45), Inches(0.85), Inches(1.7), Inches(2.5), Inches(1.0)],
    )

    add_section_title(doc, "7. Input data")
    add_table(
        doc,
        ["No", "Data fields", "Description", "Mandatory", "Valid condition", "Example"],
        usecase["input_data"],
        [Inches(0.4), Inches(1.2), Inches(1.8), Inches(1.0), Inches(1.7), Inches(1.0)],
    )

    add_section_title(doc, "8. Output data")
    if usecase.get("output_intro"):
        add_para(doc, usecase["output_intro"])
    add_table(
        doc,
        ["No", "Data fields", "Description", "Display format", "Example"],
        usecase["output_data"],
        [Inches(0.45), Inches(1.2), Inches(2.2), Inches(1.35), Inches(1.3)],
    )

    add_section_title(doc, "9. Postconditions")
    for item in usecase["postconditions"]:
        add_para(doc, item)

    out_path = OUT_DIR / usecase["filename"]
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


USECASES = [
    {
        "code": "UC01",
        "title": "Đăng ký, đăng nhập",
        "filename": "UC01DangKyDangNhap.docx",
        "description": "Use case này mô tả quá trình người dùng đăng ký tài khoản, đăng nhập vào hệ thống, đăng xuất và đổi mật khẩu. Hệ thống kiểm tra thông tin đầu vào, xác thực tài khoản và điều hướng người dùng theo vai trò Học viên, Gia sư hoặc Quản trị viên.",
        "actors": ["Học viên", "Gia sư", "Quản trị viên"],
        "preconditions": ["Người dùng đã truy cập vào hệ thống.", "Hệ thống kết nối được với cơ sở dữ liệu người dùng.", "Người dùng có email hợp lệ để đăng ký hoặc đăng nhập."],
        "basic_flow": ["Người dùng truy cập chức năng tài khoản.", "Người dùng chọn đăng ký hoặc đăng nhập.", "Người dùng nhập thông tin tài khoản.", "Hệ thống kiểm tra dữ liệu đầu vào.", "Hệ thống xác thực hoặc tạo mới tài khoản.", "Hệ thống tạo phiên đăng nhập nếu xử lý thành công.", "Hệ thống điều hướng người dùng đến trang phù hợp với vai trò."],
        "alternative_flows": [["1", "Bước 3", "Thiếu thông tin bắt buộc", "Hiển thị thông báo yêu cầu nhập đầy đủ thông tin", "Bước 3"], ["2", "Bước 4", "Email không hợp lệ", "Hiển thị thông báo email không hợp lệ", "Bước 3"], ["3", "Bước 5", "Email đã tồn tại khi đăng ký", "Hiển thị thông báo email đã tồn tại", "Bước 3"], ["4", "Bước 5", "Email hoặc mật khẩu không đúng", "Hiển thị thông báo đăng nhập thất bại", "Bước 3"], ["5", "Bước 6", "Người dùng đăng xuất", "Hệ thống xóa phiên đăng nhập", "End"]],
        "input_data": [["1", "Name", "Họ tên người dùng", "Có khi đăng ký", "Không trống", "Nguyễn Văn A"], ["2", "Email", "Email tài khoản", "Có", "Đúng định dạng email", "student@gmail.com"], ["3", "Password", "Mật khẩu", "Có", "Tối thiểu 6 ký tự", "123456"], ["4", "Role", "Vai trò tài khoản", "Có khi đăng ký", "STUDENT hoặc TUTOR", "STUDENT"], ["5", "New password", "Mật khẩu mới", "Có khi đổi mật khẩu", "Khác mật khẩu cũ", "abcdef123"]],
        "output_intro": "Thông tin tài khoản / phiên đăng nhập:",
        "output_data": [["1", "User ID", "Mã định danh người dùng", "Chuỗi ký tự", "USR001"], ["2", "Name", "Họ tên người dùng", "Chuỗi ký tự", "Nguyễn Văn A"], ["3", "Email", "Email người dùng", "Chuỗi email", "student@gmail.com"], ["4", "Role", "Vai trò người dùng", "Chuỗi ký tự", "STUDENT"], ["5", "Token", "Mã xác thực phiên đăng nhập", "JWT / cookie", ""], ["6", "Redirect", "Đường dẫn điều hướng", "URL path", "/tutor/dashboard"], ["7", "Message", "Thông báo kết quả", "Chuỗi ký tự", "Đăng nhập thành công"]],
        "postconditions": ["Tài khoản mới được tạo nếu đăng ký thành công.", "Phiên đăng nhập được tạo nếu đăng nhập thành công.", "Phiên đăng nhập bị hủy nếu người dùng đăng xuất.", "Mật khẩu mới được mã hóa và lưu nếu đổi mật khẩu thành công."],
    },
    {
        "code": "UC02",
        "title": "Quản lý hồ sơ cá nhân",
        "filename": "UC02QuanLyHoSoCaNhan.docx",
        "description": "Use case này mô tả việc người dùng xem và cập nhật thông tin cá nhân như họ tên, giới tính, ảnh đại diện và mật khẩu trên hệ thống.",
        "actors": ["Học viên", "Gia sư", "Quản trị viên"],
        "preconditions": ["Người dùng đã đăng nhập vào hệ thống.", "Tài khoản người dùng tồn tại trong cơ sở dữ liệu."],
        "basic_flow": ["Người dùng truy cập trang hồ sơ cá nhân.", "Hệ thống hiển thị thông tin hiện tại của người dùng.", "Người dùng chỉnh sửa thông tin cần cập nhật.", "Người dùng nhấn lưu thay đổi.", "Hệ thống kiểm tra dữ liệu đầu vào.", "Hệ thống cập nhật thông tin vào cơ sở dữ liệu.", "Hệ thống hiển thị thông báo cập nhật thành công."],
        "alternative_flows": [["1", "Bước 3", "Tên người dùng để trống", "Hiển thị thông báo tên không được để trống", "Bước 3"], ["2", "Bước 5", "Giới tính không hợp lệ", "Hiển thị thông báo dữ liệu không hợp lệ", "Bước 3"], ["3", "Bước 4", "Người dùng upload ảnh sai định dạng", "Hiển thị thông báo upload thất bại", "Bước 3"], ["4", "Bước 6", "Lỗi hệ thống", "Hiển thị thông báo lỗi", "End"]],
        "input_data": [["1", "Name", "Họ tên người dùng", "Không", "Không trống nếu cập nhật", "Nguyễn Văn A"], ["2", "Gender", "Giới tính", "Không", "male, female, other hoặc rỗng", "male"], ["3", "Avatar", "Ảnh đại diện", "Không", "File ảnh hợp lệ hoặc URL hợp lệ", "avatar.png"], ["4", "Current password", "Mật khẩu hiện tại", "Có khi đổi mật khẩu", "Trùng mật khẩu đang dùng", "123456"], ["5", "New password", "Mật khẩu mới", "Có khi đổi mật khẩu", "Tối thiểu 6 ký tự", "abcdef123"]],
        "output_intro": "Thông tin hồ sơ cá nhân:",
        "output_data": [["1", "User ID", "Mã người dùng", "Chuỗi ký tự", "USR001"], ["2", "Name", "Họ tên", "Chuỗi ký tự", "Nguyễn Văn A"], ["3", "Email", "Email", "Chuỗi email", "student@gmail.com"], ["4", "Gender", "Giới tính", "Chuỗi ký tự", "male"], ["5", "Avatar", "Ảnh đại diện", "URL hình ảnh", "https://..."], ["6", "Message", "Thông báo kết quả", "Chuỗi ký tự", "Cập nhật thành công"]],
        "postconditions": ["Thông tin cá nhân được cập nhật trong hệ thống.", "Người dùng nhìn thấy thông tin mới sau khi cập nhật thành công."],
    },
    {
        "code": "UC03",
        "title": "Quản lý hồ sơ gia sư",
        "filename": "UC03QuanLyHoSoGiaSu.docx",
        "description": "Use case này mô tả quá trình gia sư tạo, cập nhật và nộp hồ sơ dạy học để quản trị viên xét duyệt trước khi được hiển thị công khai trên hệ thống.",
        "actors": ["Gia sư", "Quản trị viên"],
        "preconditions": ["Gia sư đã đăng nhập vào hệ thống.", "Tài khoản gia sư đã có hoặc được tạo hồ sơ gia sư ban đầu.", "Quản trị viên đã đăng nhập khi thực hiện duyệt hồ sơ."],
        "basic_flow": ["Gia sư truy cập trang quản lý hồ sơ gia sư.", "Gia sư cập nhật thông tin cá nhân, thông tin dạy học, bằng cấp, học vấn và mạng xã hội.", "Gia sư nộp hồ sơ chờ duyệt.", "Hệ thống kiểm tra các thông tin bắt buộc.", "Hệ thống chuyển trạng thái hồ sơ sang chờ duyệt.", "Quản trị viên xem danh sách hồ sơ chờ duyệt.", "Quản trị viên duyệt hoặc từ chối hồ sơ.", "Hệ thống cập nhật trạng thái hồ sơ và gửi thông báo cho gia sư."],
        "alternative_flows": [["1", "Bước 2", "Gia sư nhập thiếu thông tin", "Lưu tạm hồ sơ ở trạng thái chưa hoàn tất", "Bước 2"], ["2", "Bước 4", "Thiếu trường bắt buộc khi nộp", "Hiển thị danh sách thông tin cần bổ sung", "Bước 2"], ["3", "Bước 7", "Admin từ chối hồ sơ", "Cập nhật trạng thái REJECTED và lưu lý do", "End"], ["4", "Bước 7", "Admin duyệt hồ sơ", "Cập nhật trạng thái APPROVED", "End"]],
        "input_data": [["1", "Bio", "Giới thiệu gia sư", "Có", "Không trống", "Có 3 năm kinh nghiệm"], ["2", "Phone", "Số điện thoại", "Có", "Chuỗi hợp lệ", "0900000000"], ["3", "Subjects", "Môn dạy", "Có", "Danh sách không rỗng", "Toán"], ["4", "Price per hour", "Học phí theo giờ", "Có", "Số lớn hơn 0", "200000"], ["5", "Qualification", "Trình độ / bằng cấp", "Có", "Không trống", "Sinh viên năm 4"], ["6", "Admin note", "Ghi chú của quản trị viên", "Không", "Chuỗi ký tự", "Cần bổ sung chứng chỉ"]],
        "output_intro": "Thông tin hồ sơ gia sư:",
        "output_data": [["1", "Tutor profile ID", "Mã hồ sơ gia sư", "Chuỗi ký tự", "TP001"], ["2", "Status", "Trạng thái hồ sơ", "Chuỗi ký tự", "APPROVED"], ["3", "Subjects", "Danh sách môn dạy", "Danh sách chuỗi", "Toán, Lý"], ["4", "Price per hour", "Học phí theo giờ", "Số tiền", "200000"], ["5", "Rating", "Điểm đánh giá trung bình", "Số từ 0 đến 5", "4.8"], ["6", "Admin note", "Ghi chú duyệt hồ sơ", "Chuỗi ký tự", "Hồ sơ hợp lệ"]],
        "postconditions": ["Hồ sơ gia sư được lưu vào hệ thống.", "Nếu được duyệt, gia sư có thể xuất hiện trong danh sách tìm kiếm công khai.", "Nếu bị từ chối, gia sư nhận lý do và có thể chỉnh sửa lại hồ sơ."],
    },
    {
        "code": "UC04",
        "title": "Tìm kiếm và xem gia sư",
        "filename": "UC04TimKiemVaXemGiaSu.docx",
        "description": "Use case này mô tả việc học viên hoặc khách truy cập tìm kiếm, lọc danh sách gia sư đã được duyệt và xem thông tin chi tiết của một gia sư.",
        "actors": ["Học viên", "Khách truy cập"],
        "preconditions": ["Người dùng đã truy cập vào hệ thống.", "Hệ thống có dữ liệu gia sư đã được duyệt."],
        "basic_flow": ["Người dùng truy cập danh sách gia sư.", "Người dùng nhập hoặc chọn tiêu chí tìm kiếm.", "Người dùng nhấn tìm kiếm hoặc hệ thống tự lọc danh sách.", "Hệ thống truy vấn danh sách gia sư phù hợp.", "Hệ thống hiển thị danh sách gia sư.", "Người dùng chọn một gia sư.", "Hệ thống hiển thị thông tin chi tiết gia sư."],
        "alternative_flows": [["1", "Bước 2", "Không nhập tiêu chí", "Hiển thị tất cả gia sư đã được duyệt", "Bước 5"], ["2", "Bước 4", "Không có kết quả phù hợp", "Hiển thị thông báo không tìm thấy gia sư", "End"], ["3", "Bước 7", "Gia sư không còn được duyệt", "Hiển thị thông báo không tìm thấy gia sư", "End"]],
        "input_data": [["1", "Subject", "Môn học cần tìm", "Không", "Chuỗi hợp lệ", "Toán"], ["2", "Area", "Khu vực học", "Không", "Chuỗi hợp lệ", "Hà Nội"], ["3", "Min price", "Mức học phí thấp nhất", "Không", "Số hợp lệ", "100000"], ["4", "Max price", "Mức học phí cao nhất", "Không", "Số hợp lệ", "300000"], ["5", "Timing shift", "Ca học mong muốn", "Không", "MORNING, AFTERNOON, EVENING hoặc FLEXIBLE", "EVENING"]],
        "output_intro": "Danh sách hoặc thông tin chi tiết gia sư:",
        "output_data": [["1", "Tutor profile ID", "Mã hồ sơ gia sư", "Chuỗi ký tự", "TP001"], ["2", "Name", "Tên gia sư", "Chuỗi ký tự", "Phạm Văn A"], ["3", "Subjects", "Môn dạy", "Danh sách chuỗi", "Toán"], ["4", "Area", "Khu vực dạy", "Danh sách chuỗi", "Hà Nội"], ["5", "Price per hour", "Học phí theo giờ", "Số tiền", "200000"], ["6", "Rating", "Đánh giá trung bình", "Số từ 0 đến 5", "4.8"], ["7", "Reviews", "Danh sách đánh giá", "Danh sách", "5 sao"]],
        "postconditions": ["Người dùng xem được danh sách gia sư phù hợp.", "Người dùng có thể chuyển sang gửi yêu cầu thuê gia sư nếu đã đăng nhập."],
    },
    {
        "code": "UC05",
        "title": "Gửi yêu cầu thuê gia sư",
        "filename": "UC05GuiYeuCauThueGiaSu.docx",
        "description": "Use case này mô tả quá trình học viên gửi yêu cầu thuê gia sư, gia sư xem yêu cầu và phản hồi bằng cách từ chối hoặc tạo khóa học.",
        "actors": ["Học viên", "Gia sư"],
        "preconditions": ["Học viên đã đăng nhập vào hệ thống.", "Gia sư được chọn có hồ sơ ở trạng thái APPROVED.", "Học viên chưa có yêu cầu đang chờ xử lý với cùng gia sư."],
        "basic_flow": ["Học viên xem chi tiết gia sư.", "Học viên mở form gửi yêu cầu học.", "Học viên nhập thông tin yêu cầu.", "Học viên gửi yêu cầu.", "Hệ thống kiểm tra dữ liệu và trạng thái gia sư.", "Hệ thống tạo yêu cầu thuê gia sư ở trạng thái PENDING.", "Hệ thống gửi thông báo cho gia sư.", "Gia sư xem yêu cầu và quyết định phản hồi."],
        "alternative_flows": [["1", "Bước 5", "Gia sư không tồn tại hoặc chưa được duyệt", "Hiển thị thông báo không tìm thấy gia sư", "End"], ["2", "Bước 5", "Học viên gửi yêu cầu cho chính mình", "Hiển thị thông báo không hợp lệ", "Bước 2"], ["3", "Bước 5", "Đã có yêu cầu PENDING với gia sư này", "Hiển thị thông báo yêu cầu đang chờ xử lý", "End"], ["4", "Bước 8", "Gia sư từ chối yêu cầu", "Cập nhật trạng thái REJECTED và thông báo cho học viên", "End"], ["5", "Bước 8", "Gia sư chấp nhận yêu cầu", "Tạo khóa học từ yêu cầu", "UC06"]],
        "input_data": [["1", "Tutor profile ID", "Mã hồ sơ gia sư", "Có", "Tồn tại và APPROVED", "TP001"], ["2", "Name", "Tên người gửi yêu cầu", "Có", "Không trống", "Nguyễn Văn A"], ["3", "Email", "Email liên hệ", "Có", "Đúng định dạng email", "student@gmail.com"], ["4", "Subject", "Môn muốn học", "Có", "Không trống", "Toán"], ["5", "Message", "Nội dung yêu cầu", "Có", "Không trống", "Muốn học 2 buổi/tuần"], ["6", "Tutor note", "Ghi chú phản hồi của gia sư", "Không", "Chuỗi ký tự", "Lịch hiện tại không phù hợp"]],
        "output_intro": "Thông tin yêu cầu thuê gia sư:",
        "output_data": [["1", "Booking ID", "Mã yêu cầu", "Chuỗi ký tự", "BK001"], ["2", "Student", "Thông tin học viên", "Đối tượng", "Nguyễn Văn A"], ["3", "Tutor", "Thông tin gia sư", "Đối tượng", "Phạm Văn B"], ["4", "Subject", "Môn học", "Chuỗi ký tự", "Toán"], ["5", "Status", "Trạng thái yêu cầu", "Chuỗi ký tự", "PENDING"], ["6", "Tutor note", "Ghi chú của gia sư", "Chuỗi ký tự", "Không phù hợp lịch"]],
        "postconditions": ["Yêu cầu thuê gia sư được tạo trong hệ thống.", "Gia sư nhận được thông báo yêu cầu mới.", "Yêu cầu có thể chuyển sang ACCEPTED, REJECTED hoặc CANCELLED."],
    },
    {
        "code": "UC06",
        "title": "Quản lý khóa học",
        "filename": "UC06QuanLyKhoaHoc.docx",
        "description": "Use case này mô tả quá trình tạo khóa học từ yêu cầu thuê gia sư, quản lý lịch học, buổi học, tin nhắn, bắt đầu, hủy và kết thúc khóa học.",
        "actors": ["Học viên", "Gia sư"],
        "preconditions": ["Gia sư đã nhận được yêu cầu thuê gia sư hợp lệ.", "Học viên và gia sư đã đăng nhập vào hệ thống.", "Nếu khóa học cần học, học viên phải thanh toán để kích hoạt khóa học."],
        "basic_flow": ["Gia sư tạo khóa học từ yêu cầu thuê gia sư.", "Hệ thống tạo lịch học và các buổi học tương ứng.", "Hệ thống chuyển khóa học sang trạng thái chờ thanh toán.", "Học viên thanh toán học phí.", "Hệ thống kích hoạt khóa học.", "Gia sư bắt đầu khóa học.", "Gia sư cập nhật trạng thái buổi học.", "Học viên và gia sư xác nhận hoàn thành buổi học.", "Hai bên xác nhận kết thúc khóa học.", "Hệ thống chuyển khóa học sang hoàn thành."],
        "alternative_flows": [["1", "Bước 1", "Yêu cầu thuê gia sư không ở trạng thái PENDING", "Hiển thị thông báo không thể tạo khóa học", "End"], ["2", "Bước 2", "Thiếu lịch học", "Hiển thị thông báo yêu cầu cung cấp thời khóa biểu", "Bước 1"], ["3", "Bước 6", "Khóa học chưa được thanh toán", "Không cho phép bắt đầu khóa học", "Bước 4"], ["4", "Bước 7", "Buổi học bị hủy hoặc học viên vắng", "Cập nhật trạng thái buổi học tương ứng", "Bước 7"], ["5", "Bước 9", "Một bên chưa xác nhận kết thúc", "Gửi thông báo chờ xác nhận cho bên còn lại", "Bước 9"]],
        "input_data": [["1", "Booking request ID", "Mã yêu cầu thuê gia sư", "Có", "Booking tồn tại và PENDING", "BK001"], ["2", "Subject", "Môn học", "Có", "Không trống", "Toán"], ["3", "Start date", "Ngày bắt đầu", "Có", "Ngày hợp lệ", "01/06/2026"], ["4", "End date", "Ngày kết thúc dự kiến", "Có", "Sau ngày bắt đầu", "30/06/2026"], ["5", "Total sessions", "Tổng số buổi", "Có", "Số nguyên lớn hơn 0", "12"], ["6", "Schedules", "Lịch học lặp lại", "Có", "Có ít nhất một lịch", "T3 18:00-20:00"], ["7", "Session status", "Trạng thái buổi học", "Không", "SCHEDULED, ONGOING, COMPLETED, CANCELLED, ABSENT", "COMPLETED"]],
        "output_intro": "Thông tin khóa học:",
        "output_data": [["1", "Course ID", "Mã khóa học", "Chuỗi ký tự", "CC001"], ["2", "Status", "Trạng thái khóa học", "Chuỗi ký tự", "ONGOING"], ["3", "Schedules", "Lịch học", "Danh sách", "T3 18:00-20:00"], ["4", "Sessions", "Danh sách buổi học", "Danh sách", "Buổi 1, Buổi 2"], ["5", "Sessions done", "Số buổi đã hoàn thành", "Số nguyên", "5"], ["6", "Messages", "Tin nhắn trong khóa học", "Danh sách", "Xin chào"]],
        "postconditions": ["Khóa học được tạo và quản lý trong hệ thống.", "Trạng thái khóa học và buổi học được cập nhật theo tiến trình học.", "Khi hai bên xác nhận kết thúc, khóa học chuyển sang COMPLETED."],
    },
    {
        "code": "UC07",
        "title": "Thanh toán học phí",
        "filename": "UC07ThanhToanHocPhi.docx",
        "description": "Use case này mô tả quá trình học viên thanh toán học phí qua Stripe để kích hoạt khóa học và hệ thống ghi nhận trạng thái thanh toán.",
        "actors": ["Học viên", "Stripe", "Hệ thống"],
        "preconditions": ["Học viên đã đăng nhập vào hệ thống.", "Khóa học tồn tại và ở trạng thái PENDING_PAYMENT.", "Khóa học có thông tin học phí hợp lệ."],
        "basic_flow": ["Học viên mở chi tiết khóa học cần thanh toán.", "Học viên chọn thanh toán học phí.", "Hệ thống tạo phiên thanh toán Stripe.", "Học viên thực hiện thanh toán trên Stripe.", "Stripe trả kết quả thanh toán về hệ thống.", "Hệ thống cập nhật thanh toán sang PAID.", "Hệ thống chuyển khóa học sang UPCOMING.", "Hệ thống cộng học phí vào khoản giữ của ví gia sư.", "Hệ thống gửi thông báo cho học viên, gia sư và quản trị viên."],
        "alternative_flows": [["1", "Bước 2", "Khóa học không cần thanh toán", "Hiển thị thông báo khóa học không cần thanh toán", "End"], ["2", "Bước 3", "Không tạo được phiên Stripe", "Hiển thị thông báo lỗi thanh toán", "End"], ["3", "Bước 4", "Học viên hủy thanh toán", "Hiển thị trang thanh toán bị hủy", "End"], ["4", "Bước 5", "Phiên thanh toán hết hạn", "Cập nhật thanh toán FAILED", "End"], ["5", "Bước 5", "Thanh toán thành công nhưng webhook chậm", "Xác minh lại trạng thái bằng session ID", "Bước 6"]],
        "input_data": [["1", "Course ID", "Mã khóa học", "Có", "Khóa học tồn tại", "CC001"], ["2", "Amount", "Số tiền thanh toán", "Có", "Số lớn hơn 0", "1200000"], ["3", "Currency", "Đơn vị tiền tệ", "Có", "Mã tiền tệ hợp lệ", "usd"], ["4", "Stripe session ID", "Mã phiên Stripe", "Không", "Chuỗi hợp lệ", "cs_test_..."], ["5", "Payment intent", "Mã giao dịch Stripe", "Không", "Chuỗi hợp lệ", "pi_..."]],
        "output_intro": "Thông tin thanh toán:",
        "output_data": [["1", "Payment ID", "Mã thanh toán", "Chuỗi ký tự", "PM001"], ["2", "Payment status", "Trạng thái thanh toán", "Chuỗi ký tự", "PAID"], ["3", "Course status", "Trạng thái khóa học", "Chuỗi ký tự", "UPCOMING"], ["4", "Amount", "Số tiền", "Số tiền", "1200000"], ["5", "Stripe status", "Trạng thái từ Stripe", "Chuỗi ký tự", "paid"], ["6", "Checkout URL", "Đường dẫn thanh toán", "URL", "https://checkout.stripe.com/..."]],
        "postconditions": ["Nếu thanh toán thành công, khóa học được kích hoạt.", "Khoản tiền học phí được giữ trong ví gia sư dưới dạng escrow.", "Học viên và gia sư nhận được thông báo thanh toán."],
    },
    {
        "code": "UC08",
        "title": "Quản lý ví và rút tiền",
        "filename": "UC08QuanLyViVaRutTien.docx",
        "description": "Use case này mô tả việc gia sư xem ví, gửi yêu cầu rút tiền và quản trị viên xử lý các yêu cầu rút tiền.",
        "actors": ["Gia sư", "Quản trị viên", "Hệ thống"],
        "preconditions": ["Gia sư đã đăng nhập vào hệ thống.", "Gia sư có hồ sơ gia sư hợp lệ.", "Gia sư có số dư khả dụng khi gửi yêu cầu rút tiền."],
        "basic_flow": ["Gia sư truy cập trang ví.", "Hệ thống hiển thị số dư, tiền đang giữ và lịch sử rút tiền.", "Gia sư nhập số tiền muốn rút.", "Hệ thống kiểm tra số dư khả dụng.", "Hệ thống tạo yêu cầu rút tiền ở trạng thái PENDING.", "Hệ thống trừ số dư khả dụng của gia sư.", "Quản trị viên xem danh sách yêu cầu rút tiền.", "Quản trị viên cập nhật trạng thái xử lý.", "Hệ thống gửi thông báo kết quả cho gia sư."],
        "alternative_flows": [["1", "Bước 4", "Số tiền rút không hợp lệ", "Hiển thị thông báo số tiền không hợp lệ", "Bước 3"], ["2", "Bước 4", "Số tiền rút lớn hơn số dư", "Hiển thị thông báo số dư không đủ", "Bước 3"], ["3", "Bước 8", "Admin cập nhật FAILED", "Hoàn lại số tiền vào ví gia sư", "End"], ["4", "Bước 8", "Admin cập nhật COMPLETED", "Ghi nhận thời gian hoàn tất và thông báo cho gia sư", "End"]],
        "input_data": [["1", "Amount", "Số tiền muốn rút", "Có", "Số lớn hơn 0 và không vượt balance", "500000"], ["2", "Withdrawal status", "Trạng thái xử lý rút tiền", "Có với admin", "PROCESSING, COMPLETED hoặc FAILED", "COMPLETED"], ["3", "Wallet ID", "Mã ví gia sư", "Có", "Tồn tại trong hệ thống", "WL001"]],
        "output_intro": "Thông tin ví và yêu cầu rút tiền:",
        "output_data": [["1", "Balance", "Số dư có thể rút", "Số tiền", "500000"], ["2", "Held amount", "Số tiền đang giữ", "Số tiền", "1000000"], ["3", "Total earned", "Tổng tiền đã kiếm", "Số tiền", "1500000"], ["4", "Withdrawal ID", "Mã yêu cầu rút tiền", "Chuỗi ký tự", "WD001"], ["5", "Status", "Trạng thái rút tiền", "Chuỗi ký tự", "PENDING"], ["6", "Completed at", "Thời gian hoàn tất", "Ngày giờ", "28/05/2026"]],
        "postconditions": ["Yêu cầu rút tiền được ghi nhận trong hệ thống.", "Số dư ví được cập nhật theo trạng thái xử lý.", "Gia sư nhận được thông báo khi yêu cầu rút tiền hoàn tất hoặc thất bại."],
    },
    {
        "code": "UC09",
        "title": "Đánh giá gia sư",
        "filename": "UC09DanhGiaGiaSu.docx",
        "description": "Use case này mô tả việc học viên đánh giá gia sư sau khi khóa học hoàn thành và hệ thống cập nhật điểm đánh giá trung bình của gia sư.",
        "actors": ["Học viên", "Gia sư", "Quản trị viên"],
        "preconditions": ["Học viên đã đăng nhập vào hệ thống.", "Khóa học của học viên đã hoàn thành.", "Học viên chưa đánh giá khóa học đó trước đây."],
        "basic_flow": ["Học viên truy cập khóa học đã hoàn thành.", "Học viên chọn chức năng đánh giá.", "Học viên nhập số sao và nhận xét.", "Hệ thống kiểm tra dữ liệu đánh giá.", "Hệ thống lưu đánh giá vào cơ sở dữ liệu.", "Hệ thống tính lại điểm trung bình và tổng số đánh giá của gia sư.", "Hệ thống gửi thông báo cho gia sư và quản trị viên.", "Hệ thống hiển thị đánh giá trên hồ sơ gia sư."],
        "alternative_flows": [["1", "Bước 1", "Khóa học chưa hoàn thành", "Không cho phép đánh giá", "End"], ["2", "Bước 3", "Rating ngoài khoảng 1 đến 5", "Hiển thị thông báo rating không hợp lệ", "Bước 3"], ["3", "Bước 4", "Học viên đã đánh giá khóa học", "Hiển thị thông báo đã đánh giá", "End"], ["4", "Bước 7", "Admin xóa đánh giá", "Xóa đánh giá và tính lại rating gia sư", "End"]],
        "input_data": [["1", "Course ID", "Mã khóa học", "Có", "Khóa học tồn tại và COMPLETED", "CC001"], ["2", "Rating", "Số sao đánh giá", "Có", "Số nguyên từ 1 đến 5", "5"], ["3", "Comment", "Nội dung nhận xét", "Không", "Chuỗi ký tự", "Gia sư dạy dễ hiểu"]],
        "output_intro": "Thông tin đánh giá:",
        "output_data": [["1", "Review ID", "Mã đánh giá", "Chuỗi ký tự", "RV001"], ["2", "Rating", "Số sao đánh giá", "Số nguyên", "5"], ["3", "Comment", "Nhận xét", "Chuỗi ký tự", "Gia sư dạy tốt"], ["4", "Tutor rating", "Điểm trung bình mới của gia sư", "Số từ 0 đến 5", "4.8"], ["5", "Total reviews", "Tổng số đánh giá", "Số nguyên", "10"]],
        "postconditions": ["Đánh giá được lưu vào hệ thống.", "Điểm trung bình và tổng số đánh giá của gia sư được cập nhật.", "Gia sư và quản trị viên nhận được thông báo đánh giá mới."],
    },
    {
        "code": "UC10",
        "title": "Quản lý thông báo",
        "filename": "UC10QuanLyThongBao.docx",
        "description": "Use case này mô tả việc người dùng nhận thông báo thời gian thực, xem danh sách thông báo, đánh dấu đã đọc và xóa thông báo.",
        "actors": ["Học viên", "Gia sư", "Quản trị viên", "Hệ thống"],
        "preconditions": ["Người dùng đã đăng nhập vào hệ thống.", "Hệ thống có cơ chế tạo thông báo khi phát sinh sự kiện nghiệp vụ."],
        "basic_flow": ["Người dùng đăng nhập vào hệ thống.", "Client kết nối luồng thông báo thời gian thực.", "Hệ thống phát sinh sự kiện nghiệp vụ.", "Hệ thống tạo thông báo trong cơ sở dữ liệu.", "Hệ thống đẩy thông báo đến người dùng nếu đang online.", "Người dùng mở danh sách thông báo.", "Người dùng đánh dấu đã đọc hoặc xóa thông báo.", "Hệ thống cập nhật trạng thái thông báo."],
        "alternative_flows": [["1", "Bước 2", "Không kết nối được realtime", "Người dùng vẫn có thể xem thông báo qua danh sách", "Bước 6"], ["2", "Bước 6", "Không có thông báo", "Hiển thị danh sách rỗng", "End"], ["3", "Bước 7", "Thông báo không thuộc người dùng", "Hiển thị thông báo không tìm thấy", "End"], ["4", "Bước 7", "Người dùng chọn đánh dấu tất cả đã đọc", "Cập nhật toàn bộ thông báo chưa đọc thành đã đọc", "End"]],
        "input_data": [["1", "Notification ID", "Mã thông báo", "Có khi cập nhật/xóa", "Tồn tại và thuộc người dùng", "NT001"], ["2", "Unread only", "Lọc thông báo chưa đọc", "Không", "true hoặc false", "true"], ["3", "Page", "Trang dữ liệu", "Không", "Số nguyên lớn hơn 0", "1"], ["4", "Limit", "Số lượng mỗi trang", "Không", "Số nguyên lớn hơn 0", "20"]],
        "output_intro": "Thông tin thông báo:",
        "output_data": [["1", "Notification ID", "Mã thông báo", "Chuỗi ký tự", "NT001"], ["2", "Type", "Loại thông báo", "Chuỗi ký tự", "BOOKING_RECEIVED"], ["3", "Title", "Tiêu đề thông báo", "Chuỗi ký tự", "Yêu cầu thuê gia sư mới"], ["4", "Body", "Nội dung thông báo", "Chuỗi ký tự", "Học viên muốn học Toán"], ["5", "Is read", "Trạng thái đã đọc", "Boolean", "false"], ["6", "Unread count", "Số thông báo chưa đọc", "Số nguyên", "5"]],
        "postconditions": ["Thông báo được lưu và hiển thị cho người dùng.", "Trạng thái đọc hoặc xóa thông báo được cập nhật theo thao tác người dùng."],
    },
    {
        "code": "UC11",
        "title": "Quản trị hệ thống",
        "filename": "UC11QuanTriHeThong.docx",
        "description": "Use case này mô tả các chức năng quản trị hệ thống như xem thống kê, quản lý người dùng, hồ sơ gia sư, khóa học, thanh toán, yêu cầu rút tiền và đánh giá.",
        "actors": ["Quản trị viên"],
        "preconditions": ["Quản trị viên đã đăng nhập vào hệ thống.", "Tài khoản có vai trò ADMIN.", "Hệ thống có dữ liệu người dùng, khóa học, thanh toán hoặc đánh giá để quản lý."],
        "basic_flow": ["Quản trị viên truy cập trang quản trị.", "Hệ thống kiểm tra quyền ADMIN.", "Hệ thống hiển thị thống kê tổng quan.", "Quản trị viên chọn nhóm dữ liệu cần quản lý.", "Hệ thống hiển thị danh sách dữ liệu tương ứng.", "Quản trị viên xem chi tiết hoặc thực hiện thao tác quản trị.", "Hệ thống cập nhật dữ liệu theo thao tác.", "Hệ thống hiển thị thông báo kết quả xử lý."],
        "alternative_flows": [["1", "Bước 2", "Người dùng không phải ADMIN", "Từ chối truy cập", "End"], ["2", "Bước 5", "Không có dữ liệu", "Hiển thị danh sách rỗng", "End"], ["3", "Bước 6", "Admin khóa gia sư", "Cập nhật trạng thái hồ sơ gia sư thành SUSPENDED", "Bước 5"], ["4", "Bước 6", "Admin xử lý rút tiền thất bại", "Hoàn lại số dư vào ví gia sư", "Bước 5"], ["5", "Bước 6", "Admin xóa đánh giá", "Xóa đánh giá và tính lại rating gia sư", "Bước 5"]],
        "input_data": [["1", "Search", "Từ khóa tìm kiếm", "Không", "Chuỗi hợp lệ", "Nguyễn"], ["2", "Role", "Vai trò người dùng cần lọc", "Không", "STUDENT, TUTOR hoặc ADMIN", "TUTOR"], ["3", "Status", "Trạng thái cần lọc", "Không", "Trạng thái hợp lệ theo từng loại dữ liệu", "APPROVED"], ["4", "Page", "Trang dữ liệu", "Không", "Số nguyên lớn hơn 0", "1"], ["5", "Limit", "Số lượng mỗi trang", "Không", "Số nguyên lớn hơn 0", "20"], ["6", "Admin action", "Thao tác quản trị", "Không", "Duyệt, từ chối, khóa, xử lý, xóa", "approve"]],
        "output_intro": "Thông tin quản trị hệ thống:",
        "output_data": [["1", "Statistics", "Thống kê tổng quan hệ thống", "Đối tượng dữ liệu", "totalUsers: 100"], ["2", "Users", "Danh sách người dùng", "Danh sách", "Nguyễn Văn A"], ["3", "Courses", "Danh sách khóa học", "Danh sách", "Toán cơ bản"], ["4", "Payments", "Danh sách thanh toán", "Danh sách", "PM001"], ["5", "Withdrawals", "Danh sách yêu cầu rút tiền", "Danh sách", "WD001"], ["6", "Reviews", "Danh sách đánh giá", "Danh sách", "5 sao"], ["7", "Message", "Thông báo kết quả", "Chuỗi ký tự", "Cập nhật thành công"]],
        "postconditions": ["Dữ liệu quản trị được hiển thị hoặc cập nhật theo thao tác của quản trị viên.", "Các thay đổi quan trọng được ghi nhận và có thể phát sinh thông báo đến người dùng liên quan."],
    },
]


def main():
    created = []
    for usecase in USECASES:
        created.append(build_doc(usecase))
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
