from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TEMPLATE = Path(r"C:\DATN\Word\UseCase\UseCaseSpecification\UC1TimKiemGiaSu.docx")
OUT_PATH = Path(r"C:\DATN\Programing\DacTa_UseCase_TongQuat_HeThongGiaSu.docx")
FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(14)
TITLE_SIZE = Pt(16)


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def apply_doc_defaults(doc):
    for style_name in ("Normal", "List Paragraph", "Heading 1", "Heading 2", "Heading 3"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
            style.font.size = BODY_SIZE


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths):
    table.style = "Table Grid"
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx < len(widths):
                set_cell_width(cell, widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                if row_idx == 0 or col_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, bold=True if row_idx == 0 else None)


def add_para(doc, text="", bold=False, align=None, style=None, title=False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=TITLE_SIZE if title else BODY_SIZE, bold=bold)
    return paragraph


def add_heading(doc, text):
    add_para(doc, text, bold=True)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Paragraph")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.15)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    format_table(table, widths)
    add_para(doc)


def build_usecase(doc, uc, first=False):
    if not first:
        doc.add_page_break()

    add_para(doc, uc["title"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, title=True)
    add_heading(doc, "1. Mã use case (Use case code)")
    add_para(doc, uc["code"])

    add_heading(doc, "2. Mô tả ngắn (Brief description)")
    add_para(doc, uc["description"])

    add_heading(doc, "3. Actors")
    for idx, actor in enumerate(uc["actors"], 1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"3.{idx} Name of Actor {idx}: ")
        set_run_font(run, bold=True)
        run = paragraph.add_run(actor)
        set_run_font(run)

    add_heading(doc, "4. Preconditions")
    for item in uc["preconditions"]:
        add_bullet(doc, item)

    add_heading(doc, "5. Basic Flow of Events")
    for item in uc["basic_flow"]:
        add_bullet(doc, item)

    add_heading(doc, "6. Alternative flows")
    add_table(
        doc,
        ["No", "Location", "Condition", "Action", "Resume location"],
        uc["alternative_flows"],
        [Inches(0.39), Inches(0.83), Inches(1.63), Inches(2.36), Inches(1.28)],
    )

    add_heading(doc, "7. Input data")
    add_table(
        doc,
        ["No", "Data fields", "Description", "Mandatory", "Valid condition", "Example"],
        uc["input_data"],
        [Inches(0.37), Inches(0.9), Inches(1.77), Inches(0.98), Inches(1.48), Inches(1.18)],
    )

    add_heading(doc, "8. Output data")
    if uc.get("output_intro"):
        add_para(doc, uc["output_intro"])
    add_table(
        doc,
        ["No", "Data fields", "Description", "Display format", "Example"],
        uc["output_data"],
        [Inches(0.37), Inches(0.99), Inches(1.74), Inches(1.93), Inches(1.58)],
    )

    add_heading(doc, "9. Postconditions")
    for item in uc["postconditions"]:
        add_bullet(doc, item)


USECASES = [
    {
        "code": "UC01",
        "title": "Đăng ký, đăng nhập",
        "description": "Use case này mô tả quá trình người dùng đăng ký tài khoản, đăng nhập vào hệ thống và đăng xuất. Hệ thống xác thực thông tin, tạo phiên làm việc và điều hướng theo vai trò Học viên, Gia sư hoặc Quản trị viên.",
        "actors": ["Học viên", "Gia sư", "Quản trị viên"],
        "preconditions": ["Người dùng đã truy cập vào hệ thống.", "Hệ thống kết nối được cơ sở dữ liệu người dùng.", "Người dùng có email hợp lệ để đăng ký hoặc đăng nhập."],
        "basic_flow": ["Người dùng mở chức năng tài khoản.", "Người dùng chọn đăng ký hoặc đăng nhập.", "Người dùng nhập họ tên, email, mật khẩu và vai trò khi đăng ký; hoặc nhập email và mật khẩu khi đăng nhập.", "Hệ thống kiểm tra định dạng email, độ dài mật khẩu và vai trò.", "Hệ thống tạo tài khoản mới hoặc xác thực tài khoản hiện có.", "Hệ thống tạo token/cookie phiên đăng nhập khi xử lý thành công.", "Hệ thống điều hướng người dùng đến trang phù hợp với vai trò."],
        "alternative_flows": [["1", "Bước 3", "Thiếu thông tin bắt buộc", "Hiển thị thông báo yêu cầu nhập đầy đủ dữ liệu", "Bước 3"], ["2", "Bước 4", "Email không hợp lệ hoặc mật khẩu dưới 6 ký tự", "Hiển thị thông báo dữ liệu không hợp lệ", "Bước 3"], ["3", "Bước 5", "Email đã tồn tại khi đăng ký", "Từ chối tạo tài khoản và hiển thị lỗi", "Bước 3"], ["4", "Bước 5", "Email hoặc mật khẩu không đúng khi đăng nhập", "Hiển thị thông báo đăng nhập thất bại", "Bước 3"], ["5", "Bước 6", "Người dùng đăng xuất", "Hệ thống xóa token/cookie phiên đăng nhập", "End"]],
        "input_data": [["1", "Name", "Họ tên người dùng", "Có khi đăng ký", "Không trống", "Nguyễn Văn A"], ["2", "Email", "Email tài khoản", "Có", "Đúng định dạng email", "student@gmail.com"], ["3", "Password", "Mật khẩu", "Có", "Tối thiểu 6 ký tự", "123456"], ["4", "Role", "Vai trò tài khoản", "Có khi đăng ký", "STUDENT, TUTOR hoặc ADMIN", "STUDENT"], ["5", "Avatar", "Ảnh đại diện dạng URL", "Không", "URL hợp lệ nếu có", "https://..."]],
        "output_intro": "Thông tin tài khoản / phiên đăng nhập:",
        "output_data": [["1", "User ID", "Mã định danh người dùng", "Chuỗi UUID", "USR001"], ["2", "Name", "Họ tên người dùng", "Chuỗi ký tự", "Nguyễn Văn A"], ["3", "Email", "Email người dùng", "Chuỗi email", "student@gmail.com"], ["4", "Role", "Vai trò người dùng", "Chuỗi ký tự", "STUDENT"], ["5", "Token", "Mã xác thực phiên đăng nhập", "JWT / cookie", ""], ["6", "Message", "Thông báo kết quả", "Chuỗi ký tự", "Đăng nhập thành công"]],
        "postconditions": ["Tài khoản mới được lưu nếu đăng ký thành công.", "Phiên đăng nhập được tạo nếu đăng nhập thành công.", "Phiên đăng nhập bị hủy nếu người dùng đăng xuất."],
    },
    {
        "code": "UC02",
        "title": "Quản lý hồ sơ cá nhân",
        "description": "Use case này mô tả việc học viên hoặc gia sư xem, cập nhật thông tin cá nhân, đổi mật khẩu và tải ảnh đại diện trên hệ thống.",
        "actors": ["Học viên", "Gia sư"],
        "preconditions": ["Người dùng đã đăng nhập vào hệ thống.", "Tài khoản người dùng tồn tại trong cơ sở dữ liệu."],
        "basic_flow": ["Người dùng truy cập trang hồ sơ cá nhân.", "Hệ thống hiển thị thông tin hiện tại gồm họ tên, email, giới tính và ảnh đại diện.", "Người dùng chỉnh sửa thông tin cần cập nhật hoặc chọn đổi mật khẩu.", "Người dùng lưu thay đổi hoặc tải ảnh đại diện mới.", "Hệ thống kiểm tra dữ liệu đầu vào và quyền của người dùng.", "Hệ thống cập nhật thông tin vào cơ sở dữ liệu hoặc Cloudinary đối với ảnh đại diện.", "Hệ thống hiển thị thông báo cập nhật thành công."],
        "alternative_flows": [["1", "Bước 3", "Tên người dùng để trống", "Hiển thị thông báo tên không được để trống", "Bước 3"], ["2", "Bước 3", "Mật khẩu hiện tại không đúng", "Từ chối đổi mật khẩu và hiển thị lỗi", "Bước 3"], ["3", "Bước 4", "Upload ảnh lỗi hoặc sai định dạng", "Hiển thị thông báo upload thất bại", "Bước 3"], ["4", "Bước 5", "Phiên đăng nhập không hợp lệ", "Yêu cầu người dùng đăng nhập lại", "End"]],
        "input_data": [["1", "Name", "Họ tên người dùng", "Không", "Không trống nếu cập nhật", "Nguyễn Văn A"], ["2", "Gender", "Giới tính", "Không", "Chuỗi hợp lệ hoặc rỗng", "male"], ["3", "Avatar", "Tệp ảnh đại diện", "Không", "File ảnh hợp lệ", "avatar.png"], ["4", "Current password", "Mật khẩu hiện tại", "Có khi đổi mật khẩu", "Trùng mật khẩu đang dùng", "123456"], ["5", "New password", "Mật khẩu mới", "Có khi đổi mật khẩu", "Tối thiểu 6 ký tự", "abcdef123"]],
        "output_intro": "Thông tin hồ sơ cá nhân:",
        "output_data": [["1", "User ID", "Mã người dùng", "Chuỗi UUID", "USR001"], ["2", "Name", "Họ tên", "Chuỗi ký tự", "Nguyễn Văn A"], ["3", "Email", "Email", "Chuỗi email", "student@gmail.com"], ["4", "Gender", "Giới tính", "Chuỗi ký tự", "male"], ["5", "Avatar", "Ảnh đại diện", "URL hình ảnh", "https://..."], ["6", "Message", "Thông báo kết quả", "Chuỗi ký tự", "Cập nhật thành công"]],
        "postconditions": ["Thông tin cá nhân được cập nhật trong hệ thống.", "Mật khẩu mới được mã hóa và lưu nếu đổi mật khẩu thành công.", "Người dùng nhìn thấy thông tin mới sau khi cập nhật."],
    },
    {
        "code": "UC03",
        "title": "Quản lý hồ sơ gia sư",
        "description": "Use case này mô tả quá trình gia sư tạo, cập nhật và nộp hồ sơ dạy học để quản trị viên xét duyệt trước khi hồ sơ được hiển thị công khai.",
        "actors": ["Gia sư", "Quản trị viên"],
        "preconditions": ["Gia sư đã đăng nhập với vai trò TUTOR.", "Hồ sơ gia sư tồn tại hoặc được hệ thống tạo khi gia sư truy cập lần đầu.", "Quản trị viên đã đăng nhập khi thực hiện duyệt hồ sơ."],
        "basic_flow": ["Gia sư truy cập trang quản lý hồ sơ gia sư.", "Hệ thống hiển thị hồ sơ hiện tại, trạng thái duyệt, học vấn và mạng xã hội.", "Gia sư cập nhật thông tin cá nhân, môn dạy, khu vực, học phí, kinh nghiệm, bằng cấp và học vấn.", "Gia sư nộp hồ sơ chờ duyệt.", "Hệ thống kiểm tra các trường bắt buộc gồm bio, phone, subjects, pricePerHour và qualification.", "Hệ thống chuyển trạng thái hồ sơ sang REVIEWING và gửi thông báo cho Admin.", "Quản trị viên xem danh sách hồ sơ chờ duyệt.", "Quản trị viên duyệt hoặc từ chối hồ sơ kèm ghi chú.", "Hệ thống cập nhật trạng thái APPROVED hoặc REJECTED và gửi thông báo cho gia sư."],
        "alternative_flows": [["1", "Bước 3", "Hồ sơ đang REVIEWING, APPROVED hoặc SUSPENDED", "Không cho chỉnh sửa và hiển thị thông báo", "End"], ["2", "Bước 5", "Thiếu trường bắt buộc", "Hiển thị danh sách thông tin cần bổ sung", "Bước 3"], ["3", "Bước 8", "Admin từ chối hồ sơ", "Cập nhật trạng thái REJECTED và lưu adminNote", "End"], ["4", "Bước 8", "Admin duyệt hồ sơ", "Cập nhật trạng thái APPROVED và cho phép hiển thị công khai", "End"]],
        "input_data": [["1", "Bio", "Giới thiệu gia sư", "Có", "Tối thiểu 10 ký tự", "Có 3 năm kinh nghiệm"], ["2", "Phone", "Số điện thoại", "Có", "9 đến 11 chữ số", "0900000000"], ["3", "Subjects", "Danh sách môn dạy", "Có", "Có ít nhất 1 môn", "Toán"], ["4", "Price per hour", "Học phí theo giờ", "Có", "Số lớn hơn 0", "200000"], ["5", "Qualification", "Trình độ / bằng cấp", "Có", "Tối thiểu 2 ký tự", "Sinh viên năm 4"], ["6", "Admin note", "Lý do từ chối của admin", "Không", "Chuỗi ký tự", "Cần bổ sung chứng chỉ"]],
        "output_intro": "Thông tin hồ sơ gia sư:",
        "output_data": [["1", "Tutor profile ID", "Mã hồ sơ gia sư", "Chuỗi UUID", "TP001"], ["2", "Status", "Trạng thái hồ sơ", "Enum", "APPROVED"], ["3", "Subjects", "Danh sách môn dạy", "Danh sách chuỗi", "Toán, Lý"], ["4", "Price per hour", "Học phí theo giờ", "Số tiền", "200000"], ["5", "Rating", "Điểm đánh giá trung bình", "Số 0 đến 5", "4.8"], ["6", "Admin note", "Ghi chú duyệt hồ sơ", "Chuỗi ký tự", "Hồ sơ hợp lệ"]],
        "postconditions": ["Hồ sơ gia sư được lưu trong hệ thống.", "Nếu được duyệt, gia sư xuất hiện trong danh sách tìm kiếm công khai.", "Nếu bị từ chối, gia sư nhận lý do và có thể chỉnh sửa nộp lại khi trạng thái cho phép."],
    },
    {
        "code": "UC04",
        "title": "Tìm kiếm và xem gia sư",
        "description": "Use case này mô tả việc học viên hoặc khách truy cập tìm kiếm, lọc danh sách gia sư đã được duyệt và xem thông tin chi tiết của một gia sư.",
        "actors": ["Học viên", "Khách truy cập"],
        "preconditions": ["Người dùng đã truy cập vào hệ thống.", "Hệ thống có dữ liệu gia sư ở trạng thái APPROVED."],
        "basic_flow": ["Người dùng truy cập trang danh sách gia sư.", "Người dùng nhập hoặc chọn tiêu chí lọc như môn học, khu vực, học phí, hình thức dạy, ca học.", "Hệ thống truy vấn danh sách gia sư APPROVED phù hợp và sắp xếp theo đánh giá.", "Hệ thống hiển thị danh sách gia sư kèm thông tin tóm tắt.", "Người dùng chọn một gia sư để xem chi tiết.", "Hệ thống hiển thị hồ sơ công khai gồm thông tin cá nhân, môn dạy, học vấn, lịch rảnh và đánh giá."],
        "alternative_flows": [["1", "Bước 2", "Không nhập tiêu chí", "Hiển thị tất cả gia sư đã được duyệt theo phân trang", "Bước 4"], ["2", "Bước 3", "Không có kết quả phù hợp", "Hiển thị danh sách rỗng và thông báo không tìm thấy", "End"], ["3", "Bước 5", "Gia sư không tồn tại hoặc không APPROVED", "Hiển thị thông báo không tìm thấy gia sư", "End"]],
        "input_data": [["1", "Subject", "Môn học cần tìm", "Không", "Có trong danh sách môn", "Toán"], ["2", "Area", "Khu vực mong muốn", "Không", "Chuỗi hợp lệ", "Hà Nội"], ["3", "Min price", "Học phí tối thiểu", "Không", "Số >= 0", "100000"], ["4", "Max price", "Học phí tối đa", "Không", "Số >= minPrice", "300000"], ["5", "Tutoring style", "Hình thức dạy", "Không", "ONE_ON_ONE, GROUP, BOTH", "ONE_ON_ONE"], ["6", "Timing shift", "Ca học", "Không", "MORNING, AFTERNOON, EVENING, FLEXIBLE", "EVENING"]],
        "output_intro": "Danh sách / chi tiết gia sư:",
        "output_data": [["1", "Tutor profile ID", "Mã hồ sơ gia sư", "Chuỗi UUID", "TP001"], ["2", "Name", "Họ tên gia sư", "Chuỗi ký tự", "Phạm Văn A"], ["3", "Subjects", "Môn dạy", "Danh sách chuỗi", "Toán"], ["4", "Preferred areas", "Khu vực dạy", "Danh sách chuỗi", "Hà Nội"], ["5", "Price per hour", "Học phí theo giờ", "Số tiền", "200000"], ["6", "Rating", "Đánh giá trung bình", "Số 0 đến 5", "4.8"], ["7", "Reviews", "Danh sách đánh giá gần nhất", "Danh sách", "5 sao"]],
        "postconditions": ["Người dùng xem được danh sách gia sư phù hợp.", "Người dùng xem được chi tiết hồ sơ công khai của gia sư đã được duyệt."],
    },
    {
        "code": "UC05",
        "title": "Đặt lịch / gửi yêu cầu thuê gia sư",
        "description": "Use case này mô tả việc học viên gửi yêu cầu thuê gia sư và gia sư phản hồi bằng cách từ chối hoặc tạo khóa học từ yêu cầu đó.",
        "actors": ["Học viên", "Gia sư"],
        "preconditions": ["Học viên đã đăng nhập với vai trò STUDENT.", "Gia sư được chọn có hồ sơ APPROVED.", "Học viên chưa có yêu cầu PENDING với cùng gia sư."],
        "basic_flow": ["Học viên xem chi tiết gia sư.", "Học viên mở form gửi yêu cầu thuê gia sư.", "Học viên nhập họ tên, email, môn học và lời nhắn.", "Hệ thống kiểm tra gia sư tồn tại, đã duyệt và không phải chính người gửi.", "Hệ thống tạo BookingRequest với trạng thái PENDING.", "Hệ thống gửi thông báo booking mới cho gia sư.", "Gia sư xem danh sách yêu cầu gửi đến mình.", "Gia sư từ chối yêu cầu hoặc tạo khóa học từ booking.", "Nếu tạo khóa học, hệ thống cập nhật booking sang ACCEPTED và tạo CourseClass chờ thanh toán."],
        "alternative_flows": [["1", "Bước 4", "Gia sư không tồn tại hoặc chưa được duyệt", "Từ chối yêu cầu và hiển thị thông báo", "End"], ["2", "Bước 4", "Học viên gửi yêu cầu cho chính mình", "Từ chối yêu cầu", "End"], ["3", "Bước 4", "Đã có booking PENDING với gia sư này", "Hiển thị thông báo yêu cầu đang chờ xử lý", "End"], ["4", "Bước 8", "Gia sư từ chối", "Cập nhật trạng thái REJECTED và gửi thông báo cho học viên", "End"], ["5", "Bước 8", "Học viên hủy khi booking còn PENDING", "Cập nhật trạng thái CANCELLED và gửi thông báo cho gia sư", "End"]],
        "input_data": [["1", "Tutor profile ID", "Gia sư được thuê", "Có", "Tồn tại và APPROVED", "TP001"], ["2", "Name", "Tên người gửi", "Có", "Không trống", "Nguyễn Văn A"], ["3", "Email", "Email liên hệ", "Có", "Đúng định dạng email", "student@gmail.com"], ["4", "Subject", "Môn muốn học", "Có", "Chuỗi hợp lệ", "Toán"], ["5", "Message", "Nội dung yêu cầu", "Có", "Không trống", "Em muốn học buổi tối"], ["6", "Tutor note", "Ghi chú từ chối", "Không", "Chuỗi ký tự", "Lịch đã kín"]],
        "output_intro": "Thông tin yêu cầu thuê gia sư:",
        "output_data": [["1", "Booking ID", "Mã yêu cầu", "Chuỗi UUID", "BK001"], ["2", "Student", "Thông tin học viên", "Đối tượng", "Nguyễn Văn A"], ["3", "Tutor", "Thông tin gia sư", "Đối tượng", "Phạm Văn B"], ["4", "Subject", "Môn học", "Chuỗi ký tự", "Toán"], ["5", "Status", "Trạng thái booking", "Enum", "PENDING"], ["6", "Message", "Thông báo kết quả", "Chuỗi ký tự", "Gửi yêu cầu thành công"]],
        "postconditions": ["Booking mới được tạo ở trạng thái PENDING nếu hợp lệ.", "Gia sư nhận thông báo yêu cầu mới.", "Khi gia sư tạo khóa học, booking chuyển sang ACCEPTED và khóa học chuyển sang PENDING_PAYMENT."],
    },
    {
        "code": "UC06",
        "title": "Quản lý khóa học",
        "description": "Use case này mô tả việc gia sư tạo khóa học từ yêu cầu thuê, học viên và gia sư theo dõi khóa học, cập nhật buổi học, xác nhận hoàn thành, nhắn tin, hủy hoặc kết thúc khóa.",
        "actors": ["Học viên", "Gia sư"],
        "preconditions": ["Booking đang ở trạng thái PENDING và thuộc về gia sư khi tạo khóa.", "Người dùng đã đăng nhập và là học viên hoặc gia sư liên quan đến khóa học.", "Khóa học đã được thanh toán trước khi bắt đầu học."],
        "basic_flow": ["Gia sư tạo khóa học từ booking, nhập môn học, ngày bắt đầu, ngày kết thúc, số buổi, thời lượng, học phí và lịch học.", "Hệ thống tạo CourseClass trạng thái PENDING_PAYMENT, tạo lịch học và các buổi học.", "Hệ thống gửi thông báo yêu cầu thanh toán cho học viên.", "Sau khi học viên thanh toán thành công, khóa học chuyển sang UPCOMING.", "Gia sư bắt đầu khóa học, hệ thống chuyển trạng thái sang ONGOING.", "Gia sư cập nhật trạng thái từng buổi học khi cần.", "Học viên và gia sư xác nhận hoàn thành từng buổi học.", "Hai bên gửi và xem tin nhắn trong khóa học.", "Một bên yêu cầu kết thúc khóa học; khi cả hai bên xác nhận, khóa học chuyển sang COMPLETED.", "Hệ thống giải phóng học phí cho ví gia sư sau khi khóa hoàn thành."],
        "alternative_flows": [["1", "Bước 1", "Thiếu lịch học, ngày hoặc số buổi", "Không tạo khóa và hiển thị lỗi", "Bước 1"], ["2", "Bước 1", "Booking không thuộc gia sư hoặc không PENDING", "Từ chối tạo khóa học", "End"], ["3", "Bước 5", "Khóa chưa UPCOMING", "Không cho bắt đầu khóa", "End"], ["4", "Bước 7", "Buổi học chưa diễn ra hoặc khóa chưa ONGOING", "Không cho xác nhận hoàn thành", "End"], ["5", "Bước 9", "Chỉ một bên xác nhận kết thúc", "Ghi nhận xác nhận và chờ bên còn lại", "Bước 9"], ["6", "Bước 9", "Một bên hủy khóa", "Chuyển khóa sang CANCELLED và xử lý thanh toán theo trạng thái", "End"]],
        "input_data": [["1", "Booking request ID", "Yêu cầu thuê liên quan", "Có", "Booking PENDING", "BK001"], ["2", "Subject", "Môn học", "Có", "Không trống", "Toán"], ["3", "Start date", "Ngày bắt đầu", "Có", "Ngày hợp lệ", "2026-06-01"], ["4", "End date", "Ngày kết thúc", "Có", "Sau ngày bắt đầu", "2026-07-01"], ["5", "Total sessions", "Tổng số buổi", "Có", "Số nguyên > 0", "12"], ["6", "Schedules", "Lịch học lặp lại", "Có", "Có ít nhất 1 lịch", "T2 08:00-09:00"], ["7", "Message content", "Nội dung tin nhắn", "Có khi nhắn tin", "Không trống", "Hôm nay học chương 1"]],
        "output_intro": "Thông tin khóa học:",
        "output_data": [["1", "Course ID", "Mã khóa học", "Chuỗi UUID", "CR001"], ["2", "Status", "Trạng thái khóa", "Enum", "ONGOING"], ["3", "Schedules", "Thời khóa biểu", "Danh sách", "T2 08:00"], ["4", "Sessions", "Danh sách buổi học", "Danh sách", "Buổi 1"], ["5", "Sessions done", "Số buổi đã hoàn thành", "Số nguyên", "3"], ["6", "Payment", "Trạng thái thanh toán", "Enum", "PAID"], ["7", "Messages", "Tin nhắn trong khóa", "Danh sách", "Xin chào"]],
        "postconditions": ["Khóa học được tạo từ booking hợp lệ.", "Trạng thái khóa học phản ánh đúng tiến trình PENDING_PAYMENT, UPCOMING, ONGOING, COMPLETED hoặc CANCELLED.", "Khi khóa hoàn thành, học phí được giải phóng cho ví gia sư."],
    },
    {
        "code": "UC07",
        "title": "Thanh toán học phí",
        "description": "Use case này mô tả việc học viên thanh toán học phí khóa học thông qua Stripe Checkout và hệ thống xác nhận giao dịch để kích hoạt khóa học.",
        "actors": ["Học viên", "Stripe"],
        "preconditions": ["Học viên đã đăng nhập với vai trò STUDENT.", "Khóa học thuộc về học viên và đang ở trạng thái PENDING_PAYMENT.", "Khóa học có tổng học phí hợp lệ.", "Stripe Checkout và webhook được cấu hình."],
        "basic_flow": ["Học viên mở khóa học cần thanh toán.", "Học viên chọn thanh toán học phí.", "Hệ thống kiểm tra quyền truy cập, trạng thái khóa và tổng học phí.", "Hệ thống tạo Stripe Checkout Session và bản ghi Payment trạng thái PENDING.", "Hệ thống chuyển học viên đến trang thanh toán của Stripe.", "Học viên hoàn tất thanh toán trên Stripe.", "Stripe gửi webhook checkout.session.completed về hệ thống.", "Hệ thống cập nhật Payment sang PAID, lưu thông tin Stripe và chuyển khóa sang UPCOMING.", "Hệ thống cộng học phí vào heldAmount của ví gia sư và gửi thông báo thanh toán thành công.", "Học viên quay lại trang success và hệ thống xác minh lại trạng thái thanh toán."],
        "alternative_flows": [["1", "Bước 3", "Khóa không tồn tại hoặc không thuộc học viên", "Từ chối thanh toán", "End"], ["2", "Bước 3", "Khóa không ở PENDING_PAYMENT", "Hiển thị thông báo khóa không cần thanh toán", "End"], ["3", "Bước 3", "Khóa chưa có tổng học phí", "Hiển thị thông báo thiếu học phí", "End"], ["4", "Bước 6", "Học viên hủy thanh toán", "Stripe chuyển về trang cancel, Payment vẫn PENDING hoặc FAILED", "End"], ["5", "Bước 7", "Webhook hết hạn session", "Cập nhật Payment sang FAILED", "End"]],
        "input_data": [["1", "Course ID", "Khóa học cần thanh toán", "Có", "Tồn tại và thuộc học viên", "CR001"], ["2", "Amount", "Tổng học phí", "Có", "Số lớn hơn 0", "2400000"], ["3", "Currency", "Đơn vị tiền", "Có", "Mặc định usd", "usd"], ["4", "Session ID", "Mã phiên Stripe", "Có khi xác minh", "Do Stripe cung cấp", "cs_test_..."], ["5", "Webhook signature", "Chữ ký webhook", "Có với webhook", "Hợp lệ theo Stripe", "t=..."]],
        "output_intro": "Thông tin thanh toán:",
        "output_data": [["1", "Payment ID", "Mã giao dịch", "Chuỗi UUID", "PM001"], ["2", "Checkout URL", "Đường dẫn Stripe Checkout", "URL", "https://checkout.stripe.com/..."], ["3", "Payment status", "Trạng thái thanh toán", "Enum", "PAID"], ["4", "Course status", "Trạng thái khóa học", "Enum", "UPCOMING"], ["5", "Paid at", "Thời điểm thanh toán", "Datetime", "2026-06-01"], ["6", "Stripe status", "Trạng thái từ Stripe", "Chuỗi ký tự", "paid"]],
        "postconditions": ["Payment được tạo hoặc cập nhật theo trạng thái Stripe.", "Khóa học chuyển sang UPCOMING sau khi thanh toán thành công.", "Ví gia sư ghi nhận khoản heldAmount chờ giải phóng."],
    },
    {
        "code": "UC08",
        "title": "Quản lý ví và rút tiền",
        "description": "Use case này mô tả việc gia sư xem ví, theo dõi số dư có thể rút, tiền đang giữ, tổng thu nhập và gửi yêu cầu rút tiền để Admin xử lý.",
        "actors": ["Gia sư", "Quản trị viên"],
        "preconditions": ["Gia sư đã đăng nhập với vai trò TUTOR.", "Gia sư có hồ sơ gia sư trong hệ thống.", "Gia sư có số dư khả dụng trong ví khi rút tiền."],
        "basic_flow": ["Gia sư truy cập trang ví.", "Hệ thống lấy thông tin TutorWallet và danh sách giao dịch đã RELEASED.", "Hệ thống hiển thị balance, heldAmount, totalEarned và lịch sử yêu cầu rút tiền.", "Gia sư nhập số tiền muốn rút.", "Hệ thống kiểm tra số tiền lớn hơn 0 và không vượt quá balance.", "Hệ thống tạo Withdrawal trạng thái PENDING và trừ balance tương ứng.", "Hệ thống gửi thông báo yêu cầu rút tiền cho Admin.", "Admin xem danh sách withdrawal.", "Admin cập nhật trạng thái PROCESSING, COMPLETED hoặc FAILED.", "Nếu COMPLETED, hệ thống ghi thời điểm hoàn tất và thông báo cho gia sư; nếu FAILED, hệ thống hoàn tiền về balance."],
        "alternative_flows": [["1", "Bước 2", "Không tìm thấy hồ sơ gia sư", "Hiển thị thông báo không tìm thấy hồ sơ", "End"], ["2", "Bước 4", "Số tiền không hợp lệ", "Hiển thị thông báo số tiền không hợp lệ", "Bước 4"], ["3", "Bước 5", "Số tiền vượt quá balance", "Từ chối yêu cầu rút tiền", "Bước 4"], ["4", "Bước 9", "Admin cập nhật FAILED", "Hoàn số tiền rút về balance", "End"], ["5", "Bước 9", "Withdrawal đã COMPLETED", "Không cho cập nhật lại", "End"]],
        "input_data": [["1", "Amount", "Số tiền muốn rút", "Có", "Số > 0 và <= balance", "50"], ["2", "Withdrawal ID", "Mã yêu cầu rút tiền", "Có khi admin xử lý", "Tồn tại", "WD001"], ["3", "Status", "Trạng thái xử lý", "Có khi admin xử lý", "PROCESSING, COMPLETED, FAILED", "COMPLETED"], ["4", "Wallet ID", "Mã ví gia sư", "Tự động", "Tồn tại", "W001"]],
        "output_intro": "Thông tin ví và yêu cầu rút tiền:",
        "output_data": [["1", "Wallet ID", "Mã ví", "Chuỗi UUID", "W001"], ["2", "Balance", "Số dư có thể rút", "Số tiền", "100"], ["3", "Held amount", "Tiền đang giữ", "Số tiền", "200"], ["4", "Total earned", "Tổng thu nhập", "Số tiền", "300"], ["5", "Withdrawal status", "Trạng thái rút tiền", "Enum", "PENDING"], ["6", "Completed at", "Thời điểm hoàn tất", "Datetime", "2026-06-05"]],
        "postconditions": ["Yêu cầu rút tiền được lưu với trạng thái PENDING khi hợp lệ.", "Balance của gia sư được trừ khi tạo yêu cầu rút tiền.", "Admin có thể xử lý yêu cầu và hệ thống cập nhật trạng thái cuối cùng."],
    },
    {
        "code": "UC09",
        "title": "Đánh giá gia sư",
        "description": "Use case này mô tả việc học viên đánh giá gia sư sau khi khóa học đã hoàn thành và Admin quản lý các đánh giá trong hệ thống.",
        "actors": ["Học viên", "Gia sư", "Quản trị viên"],
        "preconditions": ["Học viên đã đăng nhập với vai trò STUDENT.", "Khóa học thuộc về học viên và đã COMPLETED.", "Khóa học chưa có đánh giá trước đó."],
        "basic_flow": ["Học viên mở chi tiết khóa học đã hoàn thành.", "Học viên nhập điểm đánh giá từ 1 đến 5 và bình luận nếu có.", "Hệ thống kiểm tra quyền học viên, trạng thái khóa và việc đã đánh giá hay chưa.", "Hệ thống tạo Review cho khóa học.", "Hệ thống tính lại rating trung bình và totalReviews của hồ sơ gia sư.", "Hệ thống gửi thông báo đánh giá mới cho gia sư/Admin.", "Admin có thể xem danh sách đánh giá, lọc đánh giá thấp và xóa đánh giá không phù hợp.", "Nếu Admin xóa đánh giá, hệ thống tính lại rating của gia sư."],
        "alternative_flows": [["1", "Bước 3", "Người dùng không phải học viên của khóa", "Từ chối đánh giá", "End"], ["2", "Bước 3", "Khóa học chưa COMPLETED", "Hiển thị thông báo chỉ đánh giá khóa đã hoàn thành", "End"], ["3", "Bước 3", "Khóa học đã có đánh giá", "Từ chối tạo đánh giá trùng", "End"], ["4", "Bước 3", "Rating ngoài khoảng 1 đến 5", "Hiển thị thông báo rating không hợp lệ", "Bước 2"], ["5", "Bước 7", "Admin xóa đánh giá", "Xóa review và tính lại rating gia sư", "End"]],
        "input_data": [["1", "Course ID", "Khóa học được đánh giá", "Có", "Thuộc học viên và COMPLETED", "CR001"], ["2", "Rating", "Điểm đánh giá", "Có", "Số nguyên từ 1 đến 5", "5"], ["3", "Comment", "Nhận xét", "Không", "Chuỗi ký tự", "Gia sư dạy dễ hiểu"], ["4", "Flagged", "Bộ lọc đánh giá thấp", "Không", "true nếu rating <= 2", "true"]],
        "output_intro": "Thông tin đánh giá:",
        "output_data": [["1", "Review ID", "Mã đánh giá", "Chuỗi UUID", "RV001"], ["2", "Rating", "Điểm đánh giá", "Số nguyên", "5"], ["3", "Comment", "Nội dung nhận xét", "Chuỗi ký tự", "Dạy tốt"], ["4", "Tutor rating", "Rating mới của gia sư", "Số 0 đến 5", "4.7"], ["5", "Total reviews", "Tổng số đánh giá", "Số nguyên", "20"]],
        "postconditions": ["Đánh giá được lưu nếu hợp lệ.", "Rating trung bình và tổng số đánh giá của gia sư được cập nhật.", "Admin có thể xóa đánh giá và hệ thống tính lại rating."],
    },
    {
        "code": "UC10",
        "title": "Quản lý thông báo",
        "description": "Use case này mô tả việc người dùng nhận thông báo thời gian thực qua SSE, xem danh sách thông báo, đánh dấu đã đọc, đánh dấu tất cả đã đọc hoặc xóa thông báo.",
        "actors": ["Học viên", "Gia sư", "Quản trị viên"],
        "preconditions": ["Người dùng đã đăng nhập hoặc có token hợp lệ khi mở SSE stream.", "Hệ thống có NotificationService và bảng Notification."],
        "basic_flow": ["Người dùng đăng nhập vào hệ thống.", "Frontend mở kết nối SSE tới endpoint thông báo.", "Hệ thống xác thực token và đăng ký client nhận sự kiện.", "Khi có sự kiện như booking, thanh toán, khóa học, review hoặc rút tiền, hệ thống tạo Notification.", "Hệ thống đẩy thông báo thời gian thực tới client đang kết nối.", "Người dùng mở trang thông báo.", "Hệ thống hiển thị danh sách thông báo theo phân trang và số lượng chưa đọc.", "Người dùng đánh dấu một thông báo hoặc tất cả thông báo là đã đọc.", "Người dùng có thể xóa thông báo của chính mình."],
        "alternative_flows": [["1", "Bước 3", "Token SSE không hợp lệ", "Từ chối kết nối stream", "End"], ["2", "Bước 5", "Client ngắt kết nối", "Hệ thống xóa client khỏi danh sách nhận SSE", "End"], ["3", "Bước 7", "Người dùng lọc unreadOnly=true", "Chỉ hiển thị thông báo chưa đọc", "Bước 7"], ["4", "Bước 8", "Thông báo không thuộc người dùng", "Không cho cập nhật và trả về không tìm thấy", "End"], ["5", "Bước 9", "Thông báo không tồn tại", "Hiển thị thông báo không tìm thấy", "End"]],
        "input_data": [["1", "Token", "Mã xác thực SSE/REST", "Có", "Token hợp lệ", "jwt..."], ["2", "Page", "Trang dữ liệu", "Không", "Số nguyên > 0", "1"], ["3", "Limit", "Số thông báo mỗi trang", "Không", "Số nguyên > 0", "20"], ["4", "Unread only", "Chỉ lấy chưa đọc", "Không", "true hoặc false", "true"], ["5", "Notification ID", "Mã thông báo", "Có khi cập nhật/xóa", "Thuộc người dùng", "NT001"]],
        "output_intro": "Thông tin thông báo:",
        "output_data": [["1", "Notification ID", "Mã thông báo", "Chuỗi UUID", "NT001"], ["2", "Type", "Loại thông báo", "Enum", "PAYMENT_SUCCESS"], ["3", "Title", "Tiêu đề", "Chuỗi ký tự", "Thanh toán thành công"], ["4", "Body", "Nội dung", "Chuỗi ký tự", "Khóa học đã được kích hoạt"], ["5", "Is read", "Trạng thái đọc", "Boolean", "false"], ["6", "Unread count", "Số thông báo chưa đọc", "Số nguyên", "3"]],
        "postconditions": ["Thông báo được lưu và gửi tới đúng người dùng.", "Người dùng có thể quản lý trạng thái đọc của thông báo.", "Thông báo bị xóa không còn xuất hiện trong danh sách của người dùng."],
    },
    {
        "code": "UC11",
        "title": "Quản trị hệ thống",
        "description": "Use case này mô tả các chức năng quản trị gồm xem thống kê, quản lý người dùng, duyệt hồ sơ gia sư, quản lý khóa học, thanh toán, yêu cầu rút tiền và đánh giá.",
        "actors": ["Quản trị viên"],
        "preconditions": ["Quản trị viên đã đăng nhập với vai trò ADMIN.", "Token quản trị hợp lệ.", "Dữ liệu người dùng, khóa học, thanh toán, ví và đánh giá tồn tại trong hệ thống."],
        "basic_flow": ["Admin truy cập trang quản trị.", "Hệ thống kiểm tra quyền ADMIN cho toàn bộ route quản trị.", "Admin xem thống kê tổng quan gồm người dùng, gia sư, học viên, khóa học, doanh thu, yêu cầu duyệt và yêu cầu rút tiền.", "Admin xem, tìm kiếm, lọc danh sách người dùng và xem chi tiết người dùng.", "Admin tạm khóa hoặc mở khóa hồ sơ gia sư.", "Admin xem danh sách hồ sơ gia sư chờ duyệt.", "Admin xem chi tiết hồ sơ và duyệt hoặc từ chối hồ sơ.", "Admin xem danh sách khóa học, chi tiết khóa học, thanh toán và yêu cầu rút tiền.", "Admin xử lý yêu cầu rút tiền sang PROCESSING, COMPLETED hoặc FAILED.", "Admin xem danh sách đánh giá, lọc đánh giá thấp và xóa đánh giá vi phạm."],
        "alternative_flows": [["1", "Bước 2", "Người dùng không phải ADMIN", "Từ chối truy cập với mã 403", "End"], ["2", "Bước 5", "Admin tự khóa chính mình", "Từ chối thao tác", "End"], ["3", "Bước 7", "Hồ sơ gia sư không tồn tại", "Hiển thị thông báo không tìm thấy", "End"], ["4", "Bước 7", "Hồ sơ đã APPROVED hoặc REJECTED trước đó", "Hiển thị thông báo trạng thái không phù hợp", "End"], ["5", "Bước 9", "Trạng thái rút tiền không hợp lệ", "Từ chối cập nhật", "Bước 9"], ["6", "Bước 10", "Đánh giá không tồn tại", "Hiển thị thông báo không tìm thấy", "End"]],
        "input_data": [["1", "Role", "Bộ lọc vai trò người dùng", "Không", "STUDENT, TUTOR, ADMIN", "TUTOR"], ["2", "Search", "Từ khóa tìm kiếm", "Không", "Chuỗi ký tự", "nguyen"], ["3", "Tutor profile ID", "Hồ sơ gia sư cần duyệt", "Có khi duyệt", "Tồn tại", "TP001"], ["4", "Admin note", "Ghi chú từ chối", "Không", "Chuỗi ký tự", "Thiếu bằng cấp"], ["5", "Withdrawal status", "Trạng thái rút tiền", "Có khi xử lý", "PROCESSING, COMPLETED, FAILED", "COMPLETED"], ["6", "Flagged", "Bộ lọc đánh giá thấp", "Không", "true hoặc false", "true"]],
        "output_intro": "Thông tin quản trị:",
        "output_data": [["1", "Stats", "Số liệu tổng quan", "Đối tượng", "totalUsers=100"], ["2", "Users", "Danh sách người dùng", "Danh sách", "20 bản ghi"], ["3", "Tutor approvals", "Danh sách hồ sơ chờ duyệt", "Danh sách", "5 hồ sơ"], ["4", "Courses", "Danh sách khóa học", "Danh sách", "CR001"], ["5", "Payments", "Danh sách thanh toán", "Danh sách", "PM001"], ["6", "Withdrawals", "Danh sách yêu cầu rút tiền", "Danh sách", "WD001"], ["7", "Reviews", "Danh sách đánh giá", "Danh sách", "RV001"]],
        "postconditions": ["Admin xem và quản lý được dữ liệu vận hành của hệ thống.", "Trạng thái hồ sơ gia sư, withdrawal và đánh giá được cập nhật đúng theo thao tác admin.", "Các thay đổi quan trọng được gửi thông báo tới người dùng liên quan."],
    },
]


def main():
    doc = Document(str(TEMPLATE)) if TEMPLATE.exists() else Document()
    clear_document(doc)
    apply_doc_defaults(doc)

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for index, usecase in enumerate(USECASES):
        build_usecase(doc, usecase, first=index == 0)

    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    main()
